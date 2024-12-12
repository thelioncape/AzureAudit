import json

from docx import Document

import subprocess

document = Document()

document.add_heading("Azure Tenancy Audit", 0)

customer = input("Enter the customer name: ")

document.add_heading(customer, 1)

print("Starting login process")
print("Please log in to your Azure account")

subprocess.check_output("az login", shell=True)

subscriptions = subprocess.check_output(
    "az account list",
    shell=True)

subs = json.loads(subscriptions)
subCount=0
for sub in subs:
    subCount += 1
    # Number each subscription and output the name and ID of it
    print("{0}. {1} | {2}".format(subCount, sub["name"], sub["id"]))

print("Please enter a comma separated list of subscriptions you would like to audit.")
print("Alternatively, type \"all\" (without quotes) to select all subscriptions.")
while True:
    try:
        selection = input("Please enter your selection: ")

        if selection.lower() == "all":
            selection = range(1, subCount)
        else:
            selection = selection.split(",")
            for item in selection:
                int(item) # Ensure all items are purely integers
        break
    except ValueError:
        print("Invalid selection. Please try again.")

document.add_heading("Azure Assessment", 2)

paragraph = "This section provides a high-level overview of the Azure tenancy" +\
    " to identify the basic configuration within Azure."

document.add_paragraph(paragraph)

document.add_heading("Azure Tenancy Overview", 3)

p = document.add_paragraph(customer + " has " + str(subCount))
if subCount == 1:
    p.add_run(" subscription")
else:
    p.add_run(" subscriptions")
p.add_run(" in their Azure tenancy.")

document.add_paragraph("The subscriptions are as follows:")
for sub in subs:
    document.add_paragraph(sub["name"], style="List Number")

document.save("Test.docx")
